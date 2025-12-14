from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from io import BytesIO

def set_run_font(run, font_name: str, font_size: float):
    """设置 run 的字体，包括中文字体"""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_format(para, line_spacing_pt: float = 28, first_line_indent_char: int = 0):
    """设置段落格式"""
    pf = para.paragraph_format
    # 固定行距 28 磅
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    # 段前段后间距为 0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # 首行缩进（2字符 = 2 * 16磅 = 32磅）
    if first_line_indent_char > 0:
        pf.first_line_indent = Pt(first_line_indent_char * 16)

def export_to_word(name: str, date_range: str, weekly_work: str, next_week_plan: str) -> bytes:
    """将周小结数据导出为标准格式 Word 文档"""
    doc = Document()
    
    # ========== 设置页边距 ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    
    # ========== 标题：周小结（日期范围）==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title_para, line_spacing_pt=28)
    title_run = title_para.add_run(f"周小结（{date_range}）")
    set_run_font(title_run, "方正小标宋简体", 22)
    
    # ========== 姓名 ==========
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name_para, line_spacing_pt=28)
    name_run = name_para.add_run(name)
    set_run_font(name_run, "仿宋_GB2312", 16)
    
    # ========== 空行（姓名后，三号字体16磅）==========
    empty_para = doc.add_paragraph()
    empty_para.paragraph_format.first_line_indent = Pt(32)
    # 添加一个空的 run 并设置为三号字体（16磅）
    empty_run = empty_para.add_run("")
    set_run_font(empty_run, "仿宋_GB2312", 16)
    
    # ========== 本周工作：==========
    work_title_para = doc.add_paragraph()
    work_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(work_title_para, line_spacing_pt=28, first_line_indent_char=2)
    work_title_run = work_title_para.add_run("本周工作：")
    set_run_font(work_title_run, "黑体", 16)
    
    # 本周工作内容
    for line in weekly_work.strip().split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing_pt=28, first_line_indent_char=2)
            run = p.add_run(line)
            set_run_font(run, "仿宋_GB2312", 16)
    
    # ========== 空行（本周工作和下周计划之间，三号字体16磅，固定行距28磅）==========
    empty_para2 = doc.add_paragraph()
    set_paragraph_format(empty_para2, line_spacing_pt=28, first_line_indent_char=2)
    empty_run2 = empty_para2.add_run("")
    set_run_font(empty_run2, "仿宋_GB2312", 16)
    
    # ========== 下周计划：==========
    plan_title_para = doc.add_paragraph()
    plan_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(plan_title_para, line_spacing_pt=28, first_line_indent_char=2)
    plan_title_run = plan_title_para.add_run("下周计划：")
    set_run_font(plan_title_run, "黑体", 16)
    
    # 下周计划内容
    for line in next_week_plan.strip().split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing_pt=28, first_line_indent_char=2)
            run = p.add_run(line)
            set_run_font(run, "仿宋_GB2312", 16)
    
    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
